from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
from __DOCS_settings import merge_files, set_column_width
from _filter_data import value_conversion


def missing_work_summaries(num_students, student_info, raw_data, teacher_info, save_directory, root_directory):
    teacher_name = teacher_info["Teacher Name"]
    course_code = teacher_info["Course Code + Section"]
    course_name = teacher_info["Course Name"]
    school = teacher_info["School"]
    course_header = f"{course_code} - {course_name} \t {teacher_name}, {school}"

    for student in range (0, num_students):
        #student info
        f_name = student_info[student][0]
        l_name = student_info[student][1]
        student_number = student_info[student][2]
        student_email = student_info[student][3]
        caregiver_emial = student_info[student][4]
        student_header = f"{l_name}, {f_name} ({student_number})"

        # create save parameters
        student_file_name = f"{l_name}, {f_name} ({student_number}).docx"
        student_file_path = f"{save_directory}/{student_file_name}"

        # doc and page set up
        document = Document()
        sections = document.sections
        margin = 457200  # 0.5 inches for top, bottom, left and right
        for section in sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # header & preamble information
        header_text = f"{l_name}, {f_name} ({student_number})\t {school} {course_code} - {teacher_name}"
        document.add_heading(header_text, 2)
        document.add_heading("Missed Assessment Summary", 0)
        p = document.add_paragraph("Over the course of the semester it is normal to be behind on some work. This summary "
                                   "sheet lists the current learning activities that are currently listed as incomplete "
                                   "in your teacher's assessment tracking. It is important that we work together in "
                                   "order to identify the most essential learning activities that need to be handed in "
                                   "to demonstrate proficiency of the learning goals in this course. Activities that "
                                   "are highlighted are to be prioritized.")
        p = document.add_paragraph("")
        p.add_run("Deadline for work to be submitted to be included in the next reporting cycle:").italic = True
        p.add_run("_____________________________________.")
        p = document.add_paragraph()

        table_header = ["Learning Activity", "Look For", "Target", "Achievement"]
        table = document.add_table(rows=1, cols=4)
        table.style = 'Medium List 2 Accent 1'
        for c in range(0, 4):
            table.rows[0].cells[c].text = table_header[c]

        activity_counter = 0
        incomplete_counter = 0
        student_index = 7 + student
        for activity in range(0, len(raw_data)):
            achievement = raw_data[activity][student_index]
            task_status = include_task(achievement)
            if achievement == 0 or achievement == 0.5:  #incomplete, vacation, absent, retry
                learning_activity = str(raw_data[activity][1])
                look_for = str(raw_data[activity][4])
                target = str(value_conversion(val=raw_data[activity][5],
                                              conversion_type="number_to_word"))
                print_achievement = str(value_conversion(val=achievement,
                                                         conversion_type="number_to_word"))

                cells = table.add_row().cells
                cells[0].text = learning_activity
                cells[1].text = look_for
                cells[2].text = target
                cells[3].text = print_achievement
                set_column_width(table=table)
                incomplete_counter += 1
            activity_counter +=1

        if incomplete_counter < 5:
            for r in range (0, 5-incomplete_counter):
                cells = table.add_row().cells

        document.add_paragraph("")
        p = document.add_paragraph("You have ")
        p.add_run(str(incomplete_counter)).bold = True
        p.add_run(" learning activity criteria that are currently not finished out of ")
        p.add_run(str(activity_counter)).bold = True
        p.add_run(" activities that have been assigned. Please remember that inconsistent demonstration of learning "
                  "makes it very difficult to determine an accurate grade at the end of the semester.")

        document.save(str(student_file_path))
        print("")

    convert(input_path=save_directory)
    merge_files(root_directory=root_directory,
                output_name="*** Missing Work MERGED PDF FILE ***",
                save_location=save_directory)

    print("\t\t MISSING WORK SUMMARIES CREATED SUCCESSFULLY")




def include_task(task):
    m = {"INC": "Incomplete",
         "n/a": "Not Required",
         "V": "Vacation",
         "A": "Absent",
         "RT": "Re-try"}
    for key in m:
        if key == task:
            return True
    return False
