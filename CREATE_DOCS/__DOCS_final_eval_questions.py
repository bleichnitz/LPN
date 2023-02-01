from __DOCS_settings import merge_files
from _clean_data import final_eval_question_doc
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_colour(fe_table, row, col, colour):
    cell_xml_element = fe_table.rows[row].cells[col]._tc
    table_cell_property = cell_xml_element.get_or_add_tcPr()
    shading_object = OxmlElement('w:shd')
    shading_object.set(qn('w:fill'), colour)
    table_cell_property.append(shading_object)
    return 0


def set_column_width(table):
    widths = (Inches(1.25), Inches(6.25))
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width
    return 0


def question_table(doc, proficiency_q, comprehensive_q, exemplary_q, achievement):
    table = doc.add_table(rows=3, cols=2)
    set_column_width(table=table)
    header_text = ["Proficient", "Comprehensive", "Exemplary"]
    header_colours = ["#6be375", "#8a9ac2", "#7388bd"]
    questions = [proficiency_q, comprehensive_q, exemplary_q]
    for row in range(0, len(header_text)):
        # creates header text in cells
        table.rows[row].cells[0].paragraphs[0].add_run(header_text[row])
        set_cell_colour(fe_table=table,
                        row=row,
                        col=0,
                        colour=header_colours[row])

        q_cell = table.rows[row].cells[1].paragraphs[0].add_run(questions[row])
        q_cell.italic = True
        if achievement < 3:
            set_cell_colour(fe_table=table, row=0, col=1, colour="#e6e3e3")
        elif achievement < 4:
            set_cell_colour(fe_table=table, row=1, col=1, colour="#e6e3e3")
        elif achievement <= 5:
            set_cell_colour(fe_table=table, row=2, col=1, colour="#e6e3e3")

    return table


def final_evaluation_questions(num_students, student_list, teacher_info, achievement_data, standards_data,
                               save_directory, root_directory):

    question_bank = final_eval_question_doc(root_directory)
    print(question_bank)
    print(len(question_bank))

    teacher_name = teacher_info["Teacher Name"]
    course_code = teacher_info["Course Code + Section"]
    school = teacher_info["School"]

    num_tasks_to_complete = input(f"How many of the FE prompts are required by the student two complete?     ")

    for student in range(0, num_students):
        f_name = student_list[student][0]
        l_name = student_list[student][1]
        student_num = student_list[student][2]
        file_name = f"{l_name}, {f_name} ({student_num})"

        # create save parameters
        student_file_name = f"{l_name}, {f_name} ({student_num}) -- FINAL EVALUATION QUESTIONS.docx"
        student_file_path = f"{save_directory}/{student_file_name}"

        document = Document()
        sections = document.sections
        margin = 457200  # 0.5 inches for top, bottom, left and right

        for section in sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        document.add_heading(f"{file_name}\t {school} {course_code} - {teacher_name}", 2)
        document.add_heading("FINAL EVALUATION QUESTIONS", 0)

        document.add_paragraph(
            f"For the Final Evaluation (FE), you are to complete at LEAST {num_tasks_to_complete} of the questions or "
            "prompts that are listed below. You MUST complete at them from different Overarching "
            "Learning Goal categories . The remaining prompt is entirely of your choosing. Of course, you "
            "can choose to complete more than the required number of prompts to demonstrate a more thorough and "
            "complete understanding of the course learning goals. Based on your learning achievement from "
            "throughout the semester, a question has been suggested for you (see the question/prompt that is "
            "highlighted in grey).")
        document.add_paragraph("")

        fe_question_categories = int(round(len(question_bank) / 3, 0))

        print("\n\n\n\n")
        print(f"Number of FE Question Categories: {fe_question_categories}")
        print("")
        print(file_name)
        for category in range(0, fe_question_categories):
            print(question_bank[category*3][0])
            document.add_heading(question_bank[category*3][0], 1)
            p = document.add_paragraph()
            standards_sc_code = question_bank[category*3][5]
            criteria_description = find_sc_description(sc_code=standards_sc_code,
                                                       standards_data=standards_data)
            p.add_run(criteria_description).italics = True
            document.add_paragraph()
            proficiency_question = question_bank[category*3][2]
            comprehensive_question = question_bank[category * 3 + 1][2]
            exemplary_question = question_bank[category * 3 + 2][2]
            achieve = achievement_data[student][category]

            question_table(doc=document,
                           proficiency_q=proficiency_question,
                           comprehensive_q=comprehensive_question,
                           exemplary_q=exemplary_question,
                           achievement=achieve)

        document.save(student_file_path)

    convert(input_path=save_directory)
    merge_files(root_directory=root_directory,
                output_name="--- Final Evaluation Questions MERGED PDF FILE ---",
                save_location=save_directory)


def find_sc_description(sc_code, standards_data):
    for row in standards_data:
        if sc_code == row[8]:
            return row[1]
    return "no criteria found"
