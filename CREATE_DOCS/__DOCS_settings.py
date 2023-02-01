from pathlib import Path
from docx.shared import Inches, Pt
from _filter_data import value_conversion
from PyPDF4 import PdfFileMerger, PdfFileReader


def set_column_width(table):
    widths = (Inches(2), Inches(5.5), Inches(1.25), Inches(1.25))
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width
    return 0


def section_labels(filter_type, olgs_sc_data_set):
    filter_types = ["OLG", "SC", "KTCA"]
    if filter_type == filter_types[0]:  # OLGs
        return standards_labels(olgs_sc_data_set=olgs_sc_data_set, filter_type=filter_type)
    elif filter_type == filter_types[1]:  # SC
        return standards_labels(olgs_sc_data_set=olgs_sc_data_set, filter_type=filter_type)
    elif filter_type == filter_types[2]:  # KTCA
        return ktca_labels()


def standards_labels(olgs_sc_data_set, filter_type):
    olgs_sc_data_set.pop(0)
    labels = []
    previous_label = ""
    for row in olgs_sc_data_set:
        current_label = row[0]
        if current_label != previous_label:
            labels.append(current_label)
        previous_label = current_label

    sub_categories = []
    descriptions = []
    for label in labels:
        category_descriptions = []
        category_sub_categories = []
        current_label = label
        for row in olgs_sc_data_set:
            if current_label == row[0]:
                sub_category = str(row[8]).capitalize()
                category_sub_categories.append(sub_category)
                category_descriptions.append(row[1])
        descriptions.append(category_descriptions)
        sub_categories.append(category_sub_categories)

    if filter_type == "OLG":
        sub_categories = None

    return {"Category Names": labels, "Sub-Category Names": sub_categories, "Category Descriptions": descriptions}


def ktca_labels():
    labels = ["Knowledge", "Thinking", "Communication", "Application"]

    sub_categories = None

    k1 = "I can demonstrate appropriate grade level subject specific knowledge of ideas, terminology, and theories."
    k2 = "I can demonstrate understanding of the significance, meaning and interconnectedness of ideas, terminology " \
         "and skills."
    k = [k1, k2]
    t1 = "I can use grade level planning skills to generate ideas, gather information, focus research, and organize " \
         "information."
    t2 = "I can make inferences, as well as interpret, analyze, synthesize and evaluate information appropriate for " \
         "this grade level."
    t3 = "I can use appropriate grade level creative and critical thinking skills and/or processes to process " \
         "information and demonstrate my learning."
    t = [t1, t2, t3]
    c1 = "I can clearly express ideas and information in a variety of forms with appropriate precision and " \
         "sophistication for this grade level."
    c2 = "I can use grade level forms and purposes of communication to reach a variety of audiences."
    c3 = "I can use appropriate grade level communication conventions and vocabulary to share my learning."
    c = [c1, c2, c3]
    a1 = "I can apply my knowledge and skills in familiar grade level appropriate contexts."
    a2 = "I can transfer my knowledge and skills to new contexts appropriate for this grade level."
    a3 = "I can make connections between contexts that are appropriate for this grade level."
    a = [a1, a2, a3]
    descriptions = [k, t, c, a]

    return {"Category Names": labels, "Sub-Category Names": sub_categories, "Category Descriptions": descriptions}


def header_preamble(doc, student_header, reporting_cycle, filter_type):

    document = doc
    # this is to customize the preamble based on the filter type
    goals = ""
    if filter_type == "OLG":
        goals = "Overarching Learning Goals"
    elif filter_type == "SC":
        goals = "course success criteria"
    elif filter_type == "KTCA":
        goals = "Achievement Chart categories"

    if reporting_cycle == "mid-term":
        document.add_heading(student_header, 2)
        document.add_heading("Mid-Term Learning Progression Summary", 0)
        p = document.add_paragraph(f"In the tables below you will find the learning activities "
                                   f"sorted by the {goals}. In each table you will see "
                                   f"the learning activity, the task-specific look for, the learning "
                                   f"target, and your achievement. The ")
        p.add_run("targets").italic = True
        p.add_run(" refer to the task specific goals and the ")
        p.add_run("achievement").italic = True
        p.add_run(" refers to the quality of learning you have demonstrated. At the bottom of each chart "
                  "is a summary of your overall achievement in relation to the learning goal as a whole. We will "
                  "use this to help determine your grades at mid-term.")

    elif reporting_cycle == "end-of-term":
        document.add_heading(student_header, 2)
        document.add_heading("End-of-Term Learning Progression Summary", 0)
        p = document.add_paragraph(f"In the tables below you will find the learning activities "
                                   f"sorted by the {goals}. In each table you will see "
                                   f"the learning activity, the task-specific look for, the learning "
                                   f"target, and your achievement. The ")
        p.add_run("targets").italic = True
        p.add_run(" refer to the task specific goals and the ")
        p.add_run("achievement").italic = True
        p.add_run(" refers to the quality of learning you have demonstrated. At the bottom of each table "
                  "is a summary of your overall achievement in relation to the learning goal as a whole.")

        document.add_paragraph("In addition, at the end of the document you will see a graph to help you "
                               "visualize your overall achievement. The light grey bars indicate the level at "
                               "which the learning targets were taught, and in blue you will see your learning "
                               "achievement so as to be able to compare the two. Other features in the table "
                               "are: \n\t (a) the light red zone at the bottom that you are expected to "
                               "surpass to earn the credit, \n\t (b) the light grey zone at the level three to "
                               "indicate the ministry standard, and \n\t (c) the green zone that "
                               "helps to identify your overall achievement in the course.")

        document.add_paragraph("The height of the green band is dynamic and changes depending upon your "
                               "achievement. The more consistent your learning is across all the criteria, the "
                               "thinner the band, giving you a more precise indication of your final grade. The "
                               "greater the differences between your highest achievement and the criteria you need "
                               "to work on the most, the wider the band, providing greater variance in what your "
                               "final grade might be determined as.")
    return None


def summary_table(doc, data_to_print, student_index):
    table_header = ["Learning Activity", "Look For", "Target", "Achievement"]
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Medium List 2 Accent 1'
    for c in range(0, 4):
        table.rows[0].cells[c].text = table_header[c]
    for activity in data_to_print:
        learning_activity = str(activity[1])
        look_for = str(activity[4])
        target = str(value_conversion(val=activity[5],
                                      conversion_type="number_to_word"))
        achievement = str(value_conversion(val=activity[student_index],
                                           conversion_type="number_to_word"))
        cells = table.add_row().cells
        cells[0].text = learning_activity
        cells[1].text = look_for
        cells[2].text = target
        cells[3].text = achievement
        set_column_width(table=table)

    return 0


def summary_status(document, f_name, consistency, highest_achievement):

    highest_achievement = str(value_conversion(val=highest_achievement, conversion_type="number_to_word"))

    if consistency:
        spc = f"{f_name} has consistently demonstrated a {highest_achievement} " \
              f"understanding of the knowledge and skills for this learning goal."
    else:
        spc = f"Incomplete learning activities (e.g. missing work, absences, vacation, etc.) " \
              f"result in inconsistent evidence of learning. {f_name} is required to complete " \
              f"missing work so that {highest_achievement} learning can be " \
              f"confidently determined."
    p = document.add_paragraph(spc)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    return 0


def merge_files(root_directory, output_name, save_location):
    merger = PdfFileMerger()
    # r"/Users/work/OneDrive - Peel District School Board/OneDrive Desktop/GradeBook/Final Evaluation Prompts.xlsx" is
    # a "raw string", meaning that it reads the backslashes in this instance as a backslash, and not as a special
    # programming character
    merge_directory = Path(save_location)
    files = []
    blank_page_template = Path(str(root_directory)+"/"+"_templates/"+"Blank PDF Page"+".pdf")
    # print(blank_page_template)

    for file in merge_directory.iterdir():
        if str(file.suffix).lower() == ".pdf" or str(file.suffix).upper == ".PDF":
            files.append(file)

    for file in files:
        # print(file)
        read_pdf = PdfFileReader(str(file))
        total_pages = read_pdf.numPages
        # print(f"\tNumber of Pages in PDF: {total_pages}")
        merger.append(str(file))
        if int(total_pages) % 2 != 0:
            # print("\tPage ADDED")
            merger.append(str(blank_page_template))

    merger.write(str(save_location)+"/"+str(output_name)+".pdf")
    # print(str(save_location)+"/_"+str(output_name)+".pdf")
    merger.close()

    return 0


def unpack_goal_array(goals):
    unpacked = []
    for category in goals:
        for item in category:
            unpacked.append(item)
    return unpacked
