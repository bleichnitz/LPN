from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
from __DOCS_settings import section_labels, summary_table, summary_status, merge_files, header_preamble
from __DOCS_create_graph import create_graph


def create_midterm_summaries(reporting_cycle,
                             num_students,
                             student_info,
                             filter_type,
                             standards_data,
                             achievement_data,
                             learning_targets,
                             highest_achievements,
                             achievement_status,
                             performance_consistency,
                             include_graph,
                             teacher_info,
                             save_directory,
                             root_directory):
    # teacher info
    teacher_name = teacher_info["Teacher Name"]
    course_code = teacher_info["Course Code + Section"]
    school = teacher_info["School"]

    # section labels & filter type info to be inserted into the preamble
    print(filter_type)
    print("")
    labels = section_labels(filter_type=filter_type,
                            olgs_sc_data_set=standards_data)
    category_label = labels["Category Names"]
    sub_category_label = labels["Sub-Category Names"]
    category_descriptions = labels["Category Descriptions"]

    for student in range(0, num_students):
        # student info
        f_name = student_info[student][0]
        l_name = student_info[student][1]
        student_number = student_info[student][2]
        student_header = f"{l_name}, {f_name} ({student_number})"

        print(f"\t{student_header}")

        student_high_achievement = highest_achievements[student]
        student_performance_consistency = performance_consistency[student]
        student_achievement_status = achievement_status[student]

        print(f"\t\t\tNum Categories: {len(category_label)}")
        print(f"\t\t\t{learning_targets} ({len(learning_targets)})")
        print(f"\t\t\t{student_high_achievement} ({len(student_high_achievement)})")
        print(f"\t\t\t{student_performance_consistency} ({len(student_performance_consistency)})")
        print(f"\t\t\t{student_achievement_status} ({len(student_achievement_status)})")

        # create save parameters
        student_file_name = f"{l_name}, {f_name} ({student_number}) -- {reporting_cycle} PROGRESS SUMMARY.docx"
        student_file_path = f"{save_directory}/{student_file_name}"
        # print(student_file_path)

        # doc and page set up
        document = Document()
        sections = document.sections
        margin = 457200  # 0.5 inches for top, bottom, left and right
        for section in sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # header & preamble information
        header_text = f"{l_name}, {f_name} ({student_number})\t {school} {course_code} - {teacher_name}"
        header_preamble(doc=document,
                        student_header=header_text,
                        reporting_cycle=reporting_cycle,
                        filter_type=filter_type)

        document.add_paragraph()

        counter = 0
        for category in range(0, len(category_label)):
            document.add_heading(category_label[category], 1)
            chart_data = achievement_data[category]
            if sub_category_label is None:  # for OLG / KTCA
                for description in range(0, len(category_descriptions[category])):
                    p = document.add_paragraph(f"● {category_descriptions[category][description]}")
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                document.add_paragraph()
                summary_table(doc=document,
                              data_to_print=chart_data,
                              student_index=student + 7)

                document.add_paragraph()

                summary_status(document=document,
                               f_name=f_name,
                               highest_achievement=student_high_achievement[category],
                               consistency=student_performance_consistency[category])
            else:  # for SC
                for sub_cat in range(0, len(sub_category_label[category])):
                    chart_data = achievement_data[counter]
                    p = document.add_paragraph()
                    p.add_run(sub_category_label[category][sub_cat]).bold = True
                    p.add_run(f"\t {category_descriptions[category][sub_cat]}")
                    summary_table(doc=document,
                                  data_to_print=chart_data,
                                  student_index=student+7)

                    document.add_paragraph()

                    summary_status(document=document,
                                   f_name=f_name,
                                   highest_achievement=student_high_achievement[counter],
                                   consistency=student_performance_consistency[counter])

                    # TODO: insert summary status e.g. target = ?????, highest achievement = ???, consistncy = ????,
                    # status = A/M/E
                    counter += 1
                    document.add_paragraph()

        if include_graph is True:
            if filter_type == "SC":
                chart_labels = sub_category_label
            else:
                chart_labels = category_label

            category_label = labels["Category Names"]
            sub_category_label = labels["Sub-Category Names"]
            category_descriptions = labels["Category Descriptions"]
            graph = create_graph(filter_type=filter_type,
                                 targets=learning_targets,
                                 achievement=highest_achievements[student],
                                 consistency=performance_consistency[student],
                                 goal_labels=chart_labels,
                                 f_name=f_name,
                                 l_name=l_name,
                                 student_num=student_number,
                                 save_path=save_directory)
            # TODO: somehow redirect the save into the graph folder instead of the summary folder

            document.add_paragraph("")
            document.add_picture(str(graph),
                                 width=Inches(7.5))

        document.save(str(student_file_path))
        print("")

    convert(input_path=save_directory)
    merge_files(root_directory=root_directory,
                output_name=f"--- {reporting_cycle} MERGED PDF FILE ---",
                save_location=save_directory)

    print(f"\t\t {reporting_cycle} SUMMARIES CREATED SUCCESSFULLY")

    return 0
